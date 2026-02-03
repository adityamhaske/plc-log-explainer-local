import os
import pandas as pd
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import fitz  # PyMuPDF
from docx import Document as DocxDocument

class KnowledgeBaseIngestor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", " ", ""]
        )

    def parse_pdf(self, file_path: str) -> List[Document]:
        """Parses PDF and extracts text + basic table structure."""
        docs = []
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            # Basic OCR/Table handling: get_text("text") handles most layouts
            # Future: add get_text("blocks") for more structured layout info
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": os.path.basename(file_path),
                        "page": page_num + 1,
                        "type": "pdf"
                    }
                ))
        return docs

    def parse_excel(self, file_path: str) -> List[Document]:
        """Parses Excel and converts to markdown tables for LLM reasoning."""
        docs = []
        xl = pd.ExcelFile(file_path)
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            # Convert to markdown for high-fidelity table reasoning
            md_table = df.to_markdown(index=False)
            docs.append(Document(
                page_content=f"Sheet: {sheet_name}\n\n{md_table}",
                metadata={
                    "source": os.path.basename(file_path),
                    "sheet": sheet_name,
                    "type": "excel"
                }
            ))
        return docs

    def parse_docx(self, file_path: str) -> List[Document]:
        """Parses Word/Docx files."""
        docs = []
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        content = "\n".join(full_text)
        if content.strip():
            docs.append(Document(
                page_content=content,
                metadata={
                    "source": os.path.basename(file_path),
                    "type": "docx"
                }
            ))
        return docs

    def process_file(self, file_path: str) -> List[Document]:
        """Processes a file based on its extension."""
        ext = os.path.splitext(file_path)[1].lower()
        raw_docs = []
        
        if ext == '.pdf':
            raw_docs = self.parse_pdf(file_path)
        elif ext in ['.xlsx', '.xls']:
            raw_docs = self.parse_excel(file_path)
        elif ext == '.docx':
            raw_docs = self.parse_docx(file_path)
        else:
            return []

        # Split into chunks for vector indexing
        return self.text_splitter.split_documents(raw_docs)

    def process_directory(self, directory: str) -> List[Document]:
        """Processes all supported files in a directory."""
        all_docs = []
        for root, _, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)
                all_docs.extend(self.process_file(file_path))
        return all_docs
